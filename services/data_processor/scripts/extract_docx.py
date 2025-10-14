#!/usr/bin/env python3
"""
Extract text, images, and links from DOCX files
Reads from MinIO bronze-raw bucket, processes files, and stores results
"""

import os
import sys
import logging
import re
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional
from io import BytesIO

from docx import Document
from minio import Minio
from minio.error import S3Error
import psycopg2
from psycopg2.extras import Json

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class DOCXExtractor:
    """Extract text, images, and links from DOCX files"""
    
    def __init__(self):
        # MinIO configuration
        self.minio_endpoint = os.getenv("MINIO_ENDPOINT", "minio:9000")
        self.minio_access_key = os.getenv("MINIO_ACCESS_KEY", "admin")
        self.minio_secret_key = os.getenv("MINIO_SECRET_KEY", "password123")
        
        # PostgreSQL configuration
        self.postgres_host = os.getenv("POSTGRES_HOST", "postgres")
        self.postgres_db = os.getenv("POSTGRES_DB", "pikadb")
        self.postgres_user = os.getenv("POSTGRES_USER", "pika_user")
        self.postgres_password = os.getenv("POSTGRES_PASSWORD", "pika_pass")
        
        # Initialize clients
        self.minio_client = Minio(
            self.minio_endpoint,
            access_key=self.minio_access_key,
            secret_key=self.minio_secret_key,
            secure=False
        )
        
        # Ensure buckets exist
        self.ensure_buckets()
    
    def ensure_buckets(self):
        """Create required buckets if they don't exist"""
        buckets = ["bronze-raw", "silver-processed", "images"]
        for bucket in buckets:
            if not self.minio_client.bucket_exists(bucket):
                self.minio_client.make_bucket(bucket)
                logger.info(f"✅ Created bucket: {bucket}")
    
    def get_db_connection(self):
        """Get PostgreSQL connection"""
        return psycopg2.connect(
            host=self.postgres_host,
            database=self.postgres_db,
            user=self.postgres_user,
            password=self.postgres_password
        )
    
    def extract_urls(self, text: str) -> List[Dict[str, Any]]:
        """Extract URLs from text using regex"""
        # Improved regex to avoid trailing punctuation
        url_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+[^\s<>"{}|\\^`\[\].,;:!?)]'
        
        urls = []
        for match in re.finditer(url_pattern, text):
            url = match.group().rstrip('.,;:!?)')  # Remove trailing punctuation
            start_pos = match.start()
            end_pos = match.end()
            
            # Get context (50 chars before and after)
            context_start = max(0, start_pos - 50)
            context_end = min(len(text), end_pos + 50)
            context = text[context_start:context_end].strip()
            
            # Categorize URL
            url_type = self.categorize_url(url)
            
            urls.append({
                "url": url,
                "position": start_pos,
                "context": context,
                "type": url_type
            })
        
        return urls
    
    def categorize_url(self, url: str) -> str:
        """Categorize URL based on domain and content"""
        url_lower = url.lower()
        
        if 'youtube.com' in url_lower or 'youtu.be' in url_lower:
            return 'video'
        elif 'download' in url_lower or 'drive.google' in url_lower or 'dropbox' in url_lower:
            return 'download'
        elif 'discord' in url_lower or 'forum' in url_lower or 'reddit' in url_lower:
            return 'community'
        elif 'pokemmo.com' in url_lower:
            return 'official'
        else:
            return 'external'
    
    def extract_images_from_docx(self, docx_data: bytes, doc_name: str) -> Dict[str, str]:
        """Extract images from DOCX and upload to MinIO"""
        image_map = {}
        
        try:
            # Save to temporary file with safe filename
            safe_name = re.sub(r'[^\w\-_\.]', '_', doc_name)
            temp_path = f"/tmp/{safe_name}_images.docx"
            with open(temp_path, 'wb') as f:
                f.write(docx_data)
            
            # Load document
            doc = Document(temp_path)
            
            # Extract images from document relationships
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        # Get image data
                        image_part = rel.target_part
                        image_data = image_part.blob
                        
                        # Generate unique image name
                        image_hash = hashlib.md5(image_data).hexdigest()[:12]
                        ext = rel.target_ref.split('.')[-1] if '.' in rel.target_ref else 'png'
                        image_name = f"{doc_name}_{image_hash}.{ext}"
                        
                        # Upload to MinIO
                        self.minio_client.put_object(
                            bucket_name="images",
                            object_name=image_name,
                            data=BytesIO(image_data),
                            length=len(image_data),
                            content_type="image/png"
                        )
                        
                        # Store URL
                        image_url = f"http://{self.minio_endpoint}/images/{image_name}"
                        image_map[rel.rId] = image_url
                        logger.info(f"✅ Extracted image {rel.rId} -> {image_url}")
                        
                    except Exception as e:
                        logger.error(f"❌ Failed to extract image {rel.rId}: {e}")
                        continue
            
            # Clean up temp file
            os.remove(temp_path)
            
        except Exception as e:
            logger.error(f"❌ Failed to process DOCX for images: {e}")
        
        return image_map
    
    def get_paragraph_images(self, paragraph) -> List[str]:
        """Get image rIds associated with a paragraph"""
        image_rids = []
        
        # Check for inline shapes (images)
        for run in paragraph.runs:
            if 'graphic' in run._element.xml:
                # Extract relationship IDs from XML
                xml_str = run._element.xml
                matches = re.findall(r'r:embed="([^"]+)"', xml_str)
                image_rids.extend(matches)
        
        return image_rids
    
    def chunk_text_with_metadata(
        self, 
        docx_data: bytes, 
        doc_name: str,
        chunk_size: int = 1000, 
        overlap: int = 200
    ) -> List[Dict[str, Any]]:
        """Chunk document text while tracking images and links"""
        
        # Save to temporary file with safe filename
        safe_name = re.sub(r'[^\w\-_\.]', '_', doc_name)
        temp_path = f"/tmp/{safe_name}.docx"
        
        try:
            with open(temp_path, 'wb') as f:
                f.write(docx_data)
            
            # Load document
            doc = Document(temp_path)
            
            # Extract all images first
            image_map = self.extract_images_from_docx(docx_data, doc_name)
            
            # Extract full text
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            # Extract all URLs from full text
            all_urls = self.extract_urls(full_text)
            
            # Create URL position map
            url_positions = {url['position']: url for url in all_urls}
            
            chunks = []
            current_chunk = ""
            current_images = []
            chunk_index = 0
            text_position = 0
            
            for para in doc.paragraphs:
                para_text = para.text.strip()
                
                if not para_text:
                    continue
                
                # Get images associated with this paragraph
                para_images = self.get_paragraph_images(para)
                para_image_urls = [image_map.get(rid) for rid in para_images if rid in image_map]
                
                # Check if adding this paragraph would exceed chunk size
                if len(current_chunk) + len(para_text) <= chunk_size:
                    current_chunk += para_text + "\n"
                    current_images.extend(para_image_urls)
                else:
                    # Save current chunk
                    if current_chunk:
                        # Find URLs in this chunk
                        chunk_start = text_position - len(current_chunk)
                        chunk_end = text_position
                        chunk_urls = [
                            url for pos, url in url_positions.items()
                            if chunk_start <= pos < chunk_end
                        ]
                        
                        chunks.append({
                            "chunk_index": chunk_index,
                            "content": current_chunk.strip(),
                            "images": list(set(current_images)),
                            "urls": [url['url'] for url in chunk_urls],
                            "url_details": chunk_urls,
                            "has_images": len(current_images) > 0,
                            "has_links": len(chunk_urls) > 0,
                            "image_count": len(current_images),
                            "link_count": len(chunk_urls)
                        })
                        chunk_index += 1
                    
                    # Start new chunk with overlap
                    overlap_text = current_chunk[-overlap:] if len(current_chunk) > overlap else current_chunk
                    current_chunk = overlap_text + para_text + "\n"
                    current_images = para_image_urls.copy()
                
                text_position += len(para_text) + 1  # +1 for newline
            
            # Add last chunk
            if current_chunk:
                # Find URLs in last chunk
                chunk_start = text_position - len(current_chunk)
                chunk_end = text_position
                chunk_urls = [
                    url for pos, url in url_positions.items()
                    if chunk_start <= pos < chunk_end
                ]
                
                chunks.append({
                    "chunk_index": chunk_index,
                    "content": current_chunk.strip(),
                    "images": list(set(current_images)),
                    "urls": [url['url'] for url in chunk_urls],
                    "url_details": chunk_urls,
                    "has_images": len(current_images) > 0,
                    "has_links": len(chunk_urls) > 0,
                    "image_count": len(current_images),
                    "link_count": len(chunk_urls)
                })
            
            # Clean up temp file
            os.remove(temp_path)
            
            logger.info(f"✅ Created {len(chunks)} chunks from {doc_name}")
            return chunks
            
        except Exception as e:
            logger.error(f"❌ Failed to process DOCX {doc_name}: {e}")
            if os.path.exists(temp_path):
                os.remove(temp_path)
            return []
    
    def process_docx_file(self, object_name: str) -> bool:
        """Process single DOCX file from MinIO"""
        try:
            logger.info(f"📄 Processing: {object_name}")
            
            # Download file from MinIO
            response = self.minio_client.get_object("bronze-raw", object_name)
            docx_data = response.read()
            response.close()
            response.release_conn()
            
            # Get document name
            doc_name = Path(object_name).stem
            
            # Insert document record
            conn = self.get_db_connection()
            cursor = conn.cursor()
            
            try:
                cursor.execute("""
                    INSERT INTO documents (file_name, file_path, content_type, metadata, processing_status)
                    VALUES (%s, %s, 'docx', %s, 'processing')
                    RETURNING id
                """, (
                    Path(object_name).name,
                    object_name,
                    Json({
                        "source": "bronze-raw",
                        "original_path": object_name
                    })
                ))
                
                doc_id = cursor.fetchone()[0]
                conn.commit()
                logger.info(f"✅ Inserted document with ID: {doc_id}")
                
            except Exception as e:
                conn.rollback()
                logger.error(f"❌ Failed to insert document: {e}")
                return False
            finally:
                cursor.close()
                conn.close()
            
            # Process and chunk document
            chunks = self.chunk_text_with_metadata(docx_data, doc_name)
            
            if not chunks:
                logger.warning(f"⚠️ No chunks created for {object_name}")
                return False
            
            # Insert chunks
            conn = self.get_db_connection()
            cursor = conn.cursor()
            
            try:
                for chunk in chunks:
                    chunk_id = f"{doc_name}_chunk_{chunk['chunk_index']}"
                    chunk_metadata = {
                        "chunk_index": chunk["chunk_index"],
                        "images": chunk["images"],
                        "urls": chunk["urls"],
                        "url_details": chunk["url_details"],
                        "has_images": chunk["has_images"],
                        "has_links": chunk["has_links"],
                        "image_count": chunk["image_count"],
                        "link_count": chunk["link_count"]
                    }
                    
                    cursor.execute("""
                        INSERT INTO chunks (doc_id, chunk_id, content, metadata, processing_status)
                        VALUES (%s, %s, %s, %s, 'completed')
                    """, (doc_id, chunk_id, chunk["content"], Json(chunk_metadata)))
                
                # Update document status
                cursor.execute("""
                    UPDATE documents 
                    SET processing_status = 'completed', processed_at = CURRENT_TIMESTAMP
                    WHERE id = %s
                """, (doc_id,))
                
                conn.commit()
                logger.info(f"✅ Inserted {len(chunks)} chunks for {object_name}")
                
            except Exception as e:
                conn.rollback()
                logger.error(f"❌ Failed to insert chunks: {e}")
                return False
            finally:
                cursor.close()
                conn.close()
            
            return True
            
        except Exception as e:
            logger.error(f"❌ Failed to process {object_name}: {e}")
            return False
    
    def process_all_files(self):
        """Process all DOCX files from MinIO bronze-raw bucket"""
        try:
            logger.info("🚀 Starting DOCX extraction process...")
            
            # List all DOCX files in bronze-raw bucket
            objects = self.minio_client.list_objects("bronze-raw", recursive=True)
            docx_files = [obj.object_name for obj in objects if obj.object_name.endswith('.docx')]
            
            logger.info(f"📁 Found {len(docx_files)} DOCX files to process")
            
            processed_count = 0
            failed_count = 0
            
            for object_name in docx_files:
                try:
                    if self.process_docx_file(object_name):
                        processed_count += 1
                    else:
                        failed_count += 1
                except Exception as e:
                    logger.error(f"❌ Error processing {object_name}: {e}")
                    failed_count += 1
                    continue
            
            logger.info(f"📊 Processing Summary:")
            logger.info(f"   Total files: {len(docx_files)}")
            logger.info(f"   Successfully processed: {processed_count}")
            logger.info(f"   Failed: {failed_count}")
            
            return processed_count
            
        except Exception as e:
            logger.error(f"❌ Processing failed: {e}")
            return 0

def main():
    """Main function"""
    logger.info("🔍 Starting DOCX text, image, and link extraction...")
    
    try:
        extractor = DOCXExtractor()
        result = extractor.process_all_files()
        
        if result > 0:
            logger.info(f"✅ Extraction completed! {result} files processed")
        else:
            logger.warning("⚠️ No files were processed")
            
    except Exception as e:
        logger.error(f"❌ Extraction failed: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
